import pymupdf
import io
from docx import Document
# from spire.doc import *
# from spire.doc.common import *

def extract_text_from_pdf( file_data ):
    text = ""
    doc = pymupdf.open( file_data )
    for page in doc:
        text += page.get_text() + chr(12)

    return text

def extract_text_from_docx( file_data ):
    text = ""
    doc = Document(io.BytesIO(file_data))
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

# def extract_text_from_docx( file_data ):
#     doc = Document()
#     doc.LoadFromFile( file_data )
#     text = doc.GetText()

#     with open("Output/DocumentText.txt", "w", encoding="utf-8") as file:
#         file.write( text )

#     doc.close()

#     return text